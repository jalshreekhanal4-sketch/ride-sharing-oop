#!/usr/bin/env python3
"""Build the APA7 report as a .docx, mirroring report.html."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
SCREENSHOTS = HERE.parent / "screenshots"

FONT = "Times New Roman"
SIZE = Pt(12)

doc = Document()

# ---- base document defaults ----
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = SIZE
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)

section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)


def set_font(run, bold=False, italic=False, size=SIZE):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    # force east-asian complex-script font too so Word doesn't substitute
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), FONT)
    rfonts.set(qn('w:hAnsi'), FONT)
    rfonts.set(qn('w:cs'), FONT)


def centered(text=None, bold=False, italic=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold, italic=italic)
    return p


def body_para(runs):
    """runs: list of (text, bold, italic) tuples, or plain text."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    if isinstance(runs, str):
        runs = [(runs, False, False)]
    for text, bold, italic in runs:
        r = p.add_run(text)
        set_font(r, bold=bold, italic=italic)
    return p


def heading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, bold=True)
    return p


def ref_para(runs):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    for text, bold, italic in runs:
        r = p.add_run(text)
        set_font(r, bold=bold, italic=italic)
    return p


# ===================== TITLE PAGE =====================
for _ in range(8):
    doc.add_paragraph()

centered("Object-Oriented Programming in Practice: A Ride Sharing", bold=True)
centered("System in C++ and Smalltalk", bold=True, space_after=14)
centered("Jalshree Khanal")
centered("[Institution Name]")
centered("[Course Name]")
centered("[Instructor Name]")
centered("July 15, 2026")

doc.add_page_break()

# ===================== BODY =====================
heading("Object-Oriented Programming in Practice: A Ride Sharing System in C++ and Smalltalk")
doc.add_paragraph().paragraph_format.space_after = Pt(0)

body_para([
    ("This report describes a Ride Sharing System implemented twice—once in C++ and once "
     "in GNU Smalltalk—to demonstrate three core object-oriented programming (OOP) principles: "
     "encapsulation, inheritance, and polymorphism. Both versions share an identical design: a base ", False, False),
    ("Ride", False, True),
    (" class; three subclasses, ", False, False),
    ("StandardRide", False, True),
    (", ", False, False),
    ("PremiumRide", False, True),
    (", and ", False, False),
    ("SharedRide", False, True),
    (", each pricing trips differently; a ", False, False),
    ("Driver", False, True),
    (" class that accumulates completed rides; and a ", False, False),
    ("Rider", False, True),
    (" class that accumulates requested rides. A short demonstration program stores several ride "
     "objects of different subclasses in one collection and calls fare() and rideDetails() on each "
     "polymorphically. The full source code, a build/run README, and execution screenshots are "
     "available at the project's GitHub repository: "
     "https://github.com/jalshreekhanal4-sketch/ride-sharing-oop.", False, False),
])

heading("Encapsulation")
body_para([
    ("C++ hides state using access specifiers: ", False, False), ("Ride", False, True),
    (" keeps its fields protected, and ", False, False), ("Driver", False, True),
    (" and ", False, False), ("Rider", False, True),
    (" keep their ride lists (assignedRides, requestedRides) private, exposing them only through "
     "methods such as addRide() and getDriverInfo(). Encapsulation is therefore something the "
     "programmer must actively declare (Stroustrup, 2013). Smalltalk enforces the same guarantee "
     "unconditionally: an instance variable such as assignedRides is never visible outside the "
     "object at all, since the only way to interact with a Smalltalk object is by sending it a "
     "message (Goldberg & Robson, 1983). Encapsulation is not a stylistic choice in Smalltalk; it "
     "is a property of the object model itself.", False, False),
])

heading("Inheritance")
body_para(
    "C++ expresses inheritance with the class Derived : public Base syntax, and subclass "
    "constructors explicitly forward shared state to the base constructor—for example, "
    "StandardRide's constructor delegates to Ride(rideID, pickup, dropoff, distance) in its "
    "initializer list. Smalltalk expresses the same relationship through a class-creation message "
    "sent to the superclass, Ride subclass: #StandardRide ..., reflecting Smalltalk's more dynamic "
    "style, in which even class definitions are messages sent to existing objects (Goldberg & "
    "Robson, 1983). Both implementations produce an identical single-inheritance hierarchy rooted "
    "at Ride."
)

heading("Polymorphism")
body_para(
    "C++ requires fare() and rideDetails() to be declared virtual in the base class, with fare() "
    "as a pure virtual function, enabling dynamic dispatch when objects are accessed through a "
    "std::shared_ptr<Ride> (Stroustrup, 2013). Iterating over a vector of such pointers and calling "
    "ride->fare() invokes the correct override for each concrete subclass at runtime. Smalltalk has "
    "no virtual keyword because every method call is a dynamically dispatched message by default: "
    "sending fare to any Ride causes the runtime to look up the method starting at the receiver's "
    "actual class (Free Software Foundation, 2023). StandardRide, PremiumRide, and SharedRide each "
    "simply define their own fare method, and polymorphism follows automatically."
)

heading("Discussion")
body_para(
    "The comparison highlights a difference in philosophy rather than capability. C++ is statically "
    "typed and compiled, so it makes OOP guarantees opt-in and enforces them at compile time, giving "
    "the programmer fine-grained control over which behaviors are polymorphic, encapsulated, or "
    "inherited. Smalltalk is dynamically typed and purely message-based, so the same guarantees are "
    "unconditional consequences of its object model. Despite this difference, both languages arrive "
    "at the same object-oriented design for the Ride Sharing System, suggesting that encapsulation, "
    "inheritance, and polymorphism are language-independent concepts whose enforcement mechanisms "
    "simply differ."
)

heading("Conclusion")
body_para(
    "Implementing the same Ride Sharing System in C++ and Smalltalk shows that encapsulation, "
    "inheritance, and polymorphism can be realized through different mechanisms while producing "
    "equivalent designs: C++ favors explicit, compiler-checked declarations, while Smalltalk favors "
    "implicit guarantees built into a pure message-passing object model."
)

doc.add_page_break()

# ===================== REFERENCES =====================
heading("References")
doc.add_paragraph().paragraph_format.space_after = Pt(0)
ref_para([("Free Software Foundation. (2023). ", False, False),
          ("GNU Smalltalk user's guide", False, True),
          (". https://www.gnu.org/software/smalltalk/manual-base/html_node/", False, False)])
ref_para([("Goldberg, A., & Robson, D. (1983). ", False, False),
          ("Smalltalk-80: The language and its implementation", False, True),
          (". Addison-Wesley.", False, False)])
ref_para([("Stroustrup, B. (2013). ", False, False),
          ("The C++ programming language", False, True),
          (" (4th ed.). Addison-Wesley.", False, False)])

doc.add_page_break()

# ===================== APPENDIX: SCREENSHOTS =====================
heading("Appendix: Code and Sample Output")
doc.add_paragraph().paragraph_format.space_after = Pt(0)

figures = [
    ("Figure 1", "C++ — Ride Base Class and StandardRide / PremiumRide Subclasses",
     "cpp_code.png",
     "Note. Shows encapsulated fields (protected), the pure virtual fare() declared in Ride, and "
     "each subclass's overridden fare()/rideDetails() (inheritance and polymorphism)."),
    ("Figure 2", "C++ — Driver Class Demonstrating Encapsulation", "driver_code.png",
     "Note. assignedRides is private and reachable only through addRide() and getDriverInfo(); "
     "totalEarnings() and getDriverInfo() call ride->fare() and ride->rideDetails() polymorphically."),
    ("Figure 3", "C++ — Sample Program Output", "cpp_output.png",
     "Note. Output of ./RideSharingSystem, showing polymorphic fare()/rideDetails() results for all "
     "ride types, driver earnings, and rider history."),
    ("Figure 4", "Smalltalk — Ride, StandardRide, and PremiumRide Class Definitions", "st_code.png",
     "Note. Instance variables have no accessor unless explicitly defined (encapsulation by default); "
     "Ride subclass: #StandardRide creates the subclass (inheritance); each subclass overrides fare "
     "(polymorphism)."),
    ("Figure 5", "Smalltalk — Sample Program Output", "st_output.png",
     "Note. Output of gst RideSharingSystem.st, mirroring the C++ program's behavior via message "
     "sends instead of virtual function calls."),
]

for i, (num, title, filename, note) in enumerate(figures):
    if i > 0:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(num)
    set_font(r, bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r2 = p2.add_run(title)
    set_font(r2, italic=True)

    img_path = SCREENSHOTS / filename
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run()
    run3.add_picture(str(img_path), width=Inches(6.5))

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(8)
    p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r4 = p4.add_run(note)
    set_font(r4, size=Pt(11))

out_path = HERE / "Ride_Sharing_System_Report_Khanal.docx"
doc.save(str(out_path))
print("wrote", out_path)
